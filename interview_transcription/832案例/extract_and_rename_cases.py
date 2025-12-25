#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
832案例文件提取与重命名工具

功能：
1. 扫描2023年、2024年、2025年文件夹中的案例文档
2. 排除"大推送"、"大报告"、"汇报"等非案例文件
3. 从文档中提取县名
4. 复制文件到"案例补充"文件夹，并以县名重命名

用法：
    python extract_and_rename_cases.py
    python extract_and_rename_cases.py --dry-run  # 预览模式，不实际复制文件
"""

import re
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
import json


class CountyNameExtractor:
    """县名提取器"""
    
    # 县级行政区划后缀
    COUNTY_SUFFIXES = [
        "县", "市", "区", "旗", "自治县", "自治旗",
        "县级市", "林区", "特区", "工农区"
    ]
    
    # 需要排除的关键词
    EXCLUDE_KEYWORDS = ["大推送", "大报告", "汇报"]
    
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.year_folders = ["2023年", "2024年", "2025年"]
        
    def should_exclude_file(self, filename: str) -> bool:
        """判断文件是否应该排除"""
        return any(keyword in filename for keyword in self.EXCLUDE_KEYWORDS)
    
    def extract_county_from_filename(self, filename: str) -> Optional[str]:
        """从文件名中提取县名"""
        # 移除文件扩展名
        name = filename.replace('.docx', '').replace('.doc', '')
        
        # 尝试匹配完整的省市县格式：如"江西省万安县"
        pattern = r'([\u4e00-\u9fa5]+省)?([\u4e00-\u9fa5]+市)?([\u4e00-\u9fa5]+(县|市|区|旗|自治县|自治旗))'
        match = re.search(pattern, name)
        if match:
            # 返回完整的省市县名称（如果有）
            province = match.group(1) or ""
            city = match.group(2) or ""
            county = match.group(3) or ""
            
            # 如果有省份信息，返回"省+县"格式
            if province:
                return f"{province}{county}"
            # 如果有市信息但没有省，返回"市+县"格式
            elif city:
                return f"{city}{county}"
            # 否则只返回县名
            return county
        
        return None
    
    def extract_county_from_content(self, doc_path: Path) -> Optional[str]:
        """从文档内容中提取县名"""
        try:
            doc = Document(doc_path)
            
            # 策略1：从标题提取（前3段）
            for i, para in enumerate(doc.paragraphs[:3]):
                text = para.text.strip()
                if not text:
                    continue
                    
                # 查找县名模式
                county = self._find_county_in_text(text)
                if county:
                    return county
            
            # 策略2：从"访谈团抵达XX县"提取（前20段）
            for para in doc.paragraphs[:20]:
                text = para.text.strip()
                if "访谈团" in text and ("抵达" in text or "来到" in text):
                    county = self._find_county_in_text(text)
                    if county:
                        return county
            
            # 策略3：从正文开头提取（前30段）
            for para in doc.paragraphs[:30]:
                text = para.text.strip()
                if len(text) > 20:  # 只检查较长的段落
                    county = self._find_county_in_text(text)
                    if county:
                        return county
                        
        except Exception as e:
            print(f"  ⚠️  读取文档失败: {e}")
            return None
        
        return None
    
    def _find_county_in_text(self, text: str) -> Optional[str]:
        """在文本中查找县名"""
        # 先移除一些干扰词
        text = text.replace("访谈团抵达", "").replace("访谈团来到", "")
        text = text.replace("传承队来到", "").replace("调研团抵达", "")
        
        # 模式1: 省+市+县（最完整，不包含市）
        pattern1 = r'([\u4e00-\u9fa5]+省)([\u4e00-\u9fa5]{2,}(县|自治县|市|区|旗|自治旗))'
        match = re.search(pattern1, text)
        if match:
            province = match.group(1)
            county = match.group(2)
            # 排除一些非县名的词
            if not any(word in county for word in ["农业大", "工业大", "经济", "发展"]):
                return f"{province}{county}"
        
        # 模式2: 省+市+县（完整三级）
        pattern2 = r'([\u4e00-\u9fa5]+省)([\u4e00-\u9fa5]+市)([\u4e00-\u9fa5]{2,}(县|自治县|市|区|旗|自治旗))'
        match = re.search(pattern2, text)
        if match:
            province = match.group(1)
            county = match.group(3)
            if not any(word in county for word in ["农业大", "工业大", "经济", "发展"]):
                return f"{province}{county}"
        
        # 模式3: 市+县
        pattern3 = r'([\u4e00-\u9fa5]+市)([\u4e00-\u9fa5]{2,}(县|自治县|市|区|旗|自治旗))'
        match = re.search(pattern3, text)
        if match:
            city = match.group(1)
            county = match.group(2)
            # 避免"市市"、"市县"等重复
            if city != county and not any(word in county for word in ["农业大", "工业大", "经济", "发展"]):
                return f"{city}{county}"
        
        # 模式4: 只有县名（至少2个字）
        pattern4 = r'([\u4e00-\u9fa5]{2,}(县|自治县|区|旗|自治旗))'
        match = re.search(pattern4, text)
        if match:
            county = match.group(1)
            # 排除一些明显不是县名的词
            exclude_words = ["农业大", "工业大", "经济", "发展", "上限", "下限", "突破", "传统", "现代"]
            if not any(word in county for word in exclude_words):
                return county
        
        return None
    
    def extract_county_name(self, file_path: Path) -> Tuple[Optional[str], str]:
        """
        提取县名
        
        返回: (县名, 提取方式)
        """
        filename = file_path.name
        
        # 先从文件名提取
        county = self.extract_county_from_filename(filename)
        if county:
            return county, "文件名"
        
        # 再从文档内容提取
        county = self.extract_county_from_content(file_path)
        if county:
            return county, "文档内容"
        
        return None, "未提取"
    
    def process_all_files(self, output_dir: Path, dry_run: bool = False) -> Dict:
        """
        处理所有文件
        
        参数:
            output_dir: 输出目录（案例补充文件夹）
            dry_run: 是否为预览模式（不实际复制文件）
        
        返回:
            处理报告字典
        """
        results = {
            "total": 0,
            "success": 0,
            "failed": 0,
            "excluded": 0,
            "details": []
        }
        
        # 创建输出目录
        if not dry_run:
            output_dir.mkdir(parents=True, exist_ok=True)
        
        # 遍历年份文件夹
        for year_folder in self.year_folders:
            year_path = self.base_dir / year_folder
            
            if not year_path.exists():
                print(f"⚠️  文件夹不存在: {year_folder}")
                continue
            
            print(f"\n{'='*60}")
            print(f"📂 处理: {year_folder}")
            print(f"{'='*60}")
            
            # 获取所有docx文件
            files = list(year_path.glob("*.docx")) + list(year_path.glob("*.doc"))
            files.sort()
            
            for file_path in files:
                results["total"] += 1
                filename = file_path.name
                
                # 检查是否需要排除
                if self.should_exclude_file(filename):
                    print(f"🚫 排除: {filename}")
                    results["excluded"] += 1
                    results["details"].append({
                        "原文件": filename,
                        "年份": year_folder,
                        "状态": "排除",
                        "原因": "包含排除关键词"
                    })
                    continue
                
                # 提取县名
                print(f"\n处理: {filename}")
                county_name, extract_method = self.extract_county_name(file_path)
                
                if not county_name:
                    print(f"  ❌ 未能提取县名")
                    results["failed"] += 1
                    results["details"].append({
                        "原文件": filename,
                        "年份": year_folder,
                        "状态": "失败",
                        "原因": "无法提取县名"
                    })
                    continue
                
                # 生成新文件名
                new_filename = f"{county_name}.docx"
                new_file_path = output_dir / new_filename
                
                print(f"  ✅ 县名: {county_name} ({extract_method})")
                print(f"  📝 新文件: {new_filename}")
                
                # 复制文件
                if not dry_run:
                    # 如果文件已存在，添加序号
                    counter = 1
                    while new_file_path.exists():
                        new_filename = f"{county_name}_{counter}.docx"
                        new_file_path = output_dir / new_filename
                        counter += 1
                    
                    try:
                        shutil.copy2(file_path, new_file_path)
                        print(f"  💾 已复制")
                    except Exception as e:
                        print(f"  ⚠️  复制失败: {e}")
                        results["failed"] += 1
                        results["details"].append({
                            "原文件": filename,
                            "年份": year_folder,
                            "县名": county_name,
                            "提取方式": extract_method,
                            "状态": "失败",
                            "原因": f"复制失败: {e}"
                        })
                        continue
                else:
                    print(f"  🔍 [预览模式] 将复制到: {new_filename}")
                
                results["success"] += 1
                results["details"].append({
                    "原文件": filename,
                    "年份": year_folder,
                    "县名": county_name,
                    "新文件": new_filename,
                    "提取方式": extract_method,
                    "状态": "成功"
                })
        
        return results


def generate_report(results: Dict, output_path: Path):
    """生成处理报告"""
    report_lines = [
        "=" * 80,
        "832案例文件提取与重命名报告",
        "=" * 80,
        "",
        "📊 统计摘要",
        "-" * 80,
        f"总文件数: {results['total']}",
        f"成功处理: {results['success']}",
        f"处理失败: {results['failed']}",
        f"排除文件: {results['excluded']}",
        f"成功率: {results['success']/results['total']*100:.1f}%" if results['total'] > 0 else "成功率: 0%",
        "",
        "📋 详细列表",
        "-" * 80,
    ]
    
    # 按年份分组
    by_year = {}
    for detail in results['details']:
        year = detail['年份']
        if year not in by_year:
            by_year[year] = []
        by_year[year].append(detail)
    
    for year in sorted(by_year.keys()):
        report_lines.append(f"\n【{year}】")
        for detail in by_year[year]:
            status_icon = "✅" if detail['状态'] == "成功" else "❌" if detail['状态'] == "失败" else "🚫"
            report_lines.append(f"{status_icon} {detail['原文件']}")
            if detail['状态'] == "成功":
                report_lines.append(f"   → {detail['新文件']} (提取自: {detail['提取方式']})")
            else:
                report_lines.append(f"   → {detail.get('原因', '未知原因')}")
    
    report_text = "\n".join(report_lines)
    
    # 保存报告
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(report_text)
    
    # 同时保存JSON格式
    json_path = output_path.with_suffix('.json')
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"\n📄 报告已保存:")
    print(f"  - 文本格式: {output_path}")
    print(f"  - JSON格式: {json_path}")


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="832案例文件提取与重命名工具")
    parser.add_argument("--dry-run", action="store_true", help="预览模式，不实际复制文件")
    parser.add_argument("--output", type=str, default="案例补充", help="输出文件夹名称（默认: 案例补充）")
    args = parser.parse_args()
    
    # 设置路径
    base_dir = Path(__file__).parent
    output_dir = base_dir / args.output
    
    print("=" * 80)
    print("🚀 832案例文件提取与重命名工具")
    print("=" * 80)
    print(f"📂 源目录: {base_dir}")
    print(f"📁 目标目录: {output_dir}")
    print(f"🔍 模式: {'预览模式（不复制文件）' if args.dry_run else '正式模式（将复制文件）'}")
    print()
    
    # 创建提取器并处理
    extractor = CountyNameExtractor(base_dir)
    results = extractor.process_all_files(output_dir, dry_run=args.dry_run)
    
    # 生成报告
    print(f"\n{'='*80}")
    print("📊 处理完成")
    print(f"{'='*80}")
    print(f"总文件数: {results['total']}")
    print(f"成功处理: {results['success']}")
    print(f"处理失败: {results['failed']}")
    print(f"排除文件: {results['excluded']}")
    
    if not args.dry_run:
        report_path = base_dir / "提取报告.txt"
        generate_report(results, report_path)
    else:
        print("\n💡 提示: 使用 --dry-run 选项预览结果，去掉该选项后将实际复制文件")
    
    # 如果有失败的文件，列出来
    if results['failed'] > 0:
        print(f"\n⚠️  以下文件处理失败:")
        for detail in results['details']:
            if detail['状态'] == "失败":
                print(f"  - {detail['原文件']} ({detail.get('原因', '未知原因')})")


if __name__ == "__main__":
    main()

